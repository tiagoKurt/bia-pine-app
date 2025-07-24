# bia.py
import pandas as pd
import os, re, requests
from docx import Document

def gerar_dicionario_word(recurso_url: str, template_path: str) -> str:
    """
    Gera um dicionário de dados em formato Word a partir de um recurso CKAN.
    recurso_url: link completo do recurso CKAN
    template_path: caminho do arquivo modelo_bia2_pronto_para_preencher.docx
    Retorna: caminho do arquivo Word gerado
    """

    # ====== Extrair resource_id e dataset_id ======
    match = re.search(r'/resource/([a-zA-Z0-9-]+)', recurso_url)
    if not match:
        raise ValueError('Link do recurso CKAN inválido.')
    resource_id = match.group(1)

    dataset_match = re.search(r'/dataset/([a-zA-Z0-9-]+)', recurso_url)
    if not dataset_match:
        raise ValueError('Não foi possível extrair o dataset_id do link.')
    dataset_id = dataset_match.group(1)

    # ====== Buscar informações do dataset ======
    PORTAL_URL = 'https://dadosabertos.go.gov.br'
    resp_dataset = requests.get(f"{PORTAL_URL}/api/3/action/package_show?id={dataset_id}", timeout=30)
    resp_dataset.raise_for_status()
    dataset_info = resp_dataset.json()['result']

    # Nome da base de dados (limpo)
    nome_base_raw = dataset_info.get('title','')
    nome_limpo = re.sub(
        r'\b(20\d{2}|\d{1,2}|janeiro|fevereiro|março|marco|abril|maio|junho|julho|agosto|setembro|outubro|novembro|dezembro)\b',
        '', nome_base_raw, flags=re.IGNORECASE
    )
    nome_somente_letras = ''.join(c for c in nome_limpo if c.isalpha() or c.isspace())
    nome_somente_letras = re.sub(r'\s+', ' ', nome_somente_letras).strip().upper()

    # Montar título do documento
    titulo_documento = f"DICIONÁRIO DE DADOS {nome_somente_letras}"

    # ====== Buscar descrição do recurso ======
    resp_resource = requests.get(f"{PORTAL_URL}/api/3/action/resource_show?id={resource_id}", timeout=30)
    descricao_recurso = ''
    if resp_resource.status_code == 200:
        res_info = resp_resource.json()['result']
        descricao_recurso = res_info.get('description','')
    if not descricao_recurso:
        descricao_recurso = dataset_info.get('notes','')
    if not descricao_recurso:
        descricao_recurso = f"DADOS DA BASE: {nome_somente_letras}"

    # ====== Buscar dados do recurso ======
    api_url = f"{PORTAL_URL}/api/3/action/datastore_search?resource_id={resource_id}&limit=5000"
    resp = requests.get(api_url, timeout=30)
    resp.raise_for_status()
    df_res = pd.DataFrame(resp.json()['result']['records'])

    # ====== Detectar tipo de coluna ======
    def detectar_tipo(coluna_serie):
        amostra = coluna_serie.dropna().astype(str).head(200)
        tem_letras = any(any(c.isalpha() for c in str(v)) for v in amostra)
        tem_numeros = any(any(c.isdigit() for c in str(v)) for v in amostra)
        if tem_letras and tem_numeros:
            return 'ALFANUMÉRICO'
        elif tem_numeros and not tem_letras:
            return 'NUMÉRICO'
        else:
            return 'TEXTO'

    coluna_info = []
    for coluna in df_res.columns:
        tipo_dado = detectar_tipo(df_res[coluna])
        descricao = coluna.replace('_',' ').replace('/',' ').upper()
        coluna_info.append({'coluna': coluna, 'tipo': tipo_dado, 'descricao': descricao})

    # ====== Preencher o template ======
    if not os.path.exists(template_path):
        raise FileNotFoundError('Template não encontrado.')

    doc = Document(template_path)

    for p in doc.paragraphs:
        if '{{ titulo_documento }}' in p.text:
            p.text = titulo_documento
        if '{{ descricao_base }}' in p.text:
            p.text = descricao_recurso

    # Localizar tabela e preencher
    table = None
    for t in doc.tables:
        headers = [cell.text.strip().upper() for cell in t.rows[0].cells]
        if headers == ['CAMPO/COLUNA','TIPO','DESCRIÇÃO']:
            table = t
            break
    if table is None:
        raise ValueError('Tabela padrão não encontrada no template.')

    while len(table.rows) > 1:
        table._tbl.remove(table.rows[1]._tr)

    for info in coluna_info:
        row = table.add_row().cells
        row[0].text = info['coluna']
        row[1].text = info['tipo']
        row[2].text = info['descricao']

    output_name = titulo_documento.replace(' ', '_') + '.docx'
    doc.save(output_name)
    return output_name